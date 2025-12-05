"""
AI Math Tutor - Backend Server
==============================
Student: Aric Hurkman
Course: CSCI 250
Final Project: AI Math Tutor
Professor: Gheni Abla

This Flask server provides the backend API for the AI Math Tutor application.
It handles requests to solve math problems and generate quiz questions using
the Google Gemini API with user-provided API keys.

Features:
- Google OAuth integration for user authentication
- User-provided Gemini API keys (each user uses their own key)
- Problem solving with step-by-step explanations
- Quiz generation and answer evaluation
- STUDY MODE: Interactive step-by-step learning with hints
- FILE UPLOAD SUPPORT: Images (.png, .jpg, .jpeg, .gif, .webp), PDFs, and DOCX files

Requirements:
- Flask: Web framework for Python
- Flask-CORS: Cross-Origin Resource Sharing support
- google-generativeai: Google's Python SDK for Gemini API
- Pillow: Image processing library
- pdf2image: PDF to image conversion
- python-docx: Word document text extraction
- PyMuPDF (fitz): PDF text extraction

Usage:
    python server.py
    
The server will start on http://localhost:5000
"""

# =============================================================================
# IMPORTS
# =============================================================================

# Flask framework for creating the web server
from flask import Flask, request, jsonify, send_from_directory

# CORS (Cross-Origin Resource Sharing) to allow frontend to communicate with backend
from flask_cors import CORS

# Google Generative AI SDK for Gemini API integration
import google.generativeai as genai

# OS module for environment variable access and file operations
import os

# JSON module for parsing responses
import json

# Regular expressions for cleaning JSON responses
import re

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

# Base64 encoding for file uploads
import base64

# IO module for handling byte streams
import io

# Temporary file handling
import tempfile

# PIL/Pillow for image processing
from PIL import Image

# PDF processing libraries
try:
    import fitz  # PyMuPDF for PDF text and image extraction
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# Word document processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# =============================================================================
# APPLICATION CONFIGURATION
# =============================================================================

# Initialize Flask application with static file serving
# static_folder='.' means serve static files from the current directory
# static_url_path='' means serve them from the root URL
app = Flask(__name__, static_folder='.', static_url_path='')

# Enable CORS for all routes
# This allows the frontend (running on a different port) to make requests to this server
CORS(app)

# Configure maximum file upload size (16 MB)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# Allowed file extensions for uploads
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp'}
ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'docx', 'doc'}

# =============================================================================
# SYSTEM PROMPTS
# =============================================================================

# System prompt for solving math problems
SOLVER_SYSTEM_PROMPT = """You are an expert math tutor helping students understand mathematical concepts. 
Your role is to solve math problems step-by-step with clear, educational explanations.

When solving problems:
1. First, identify what type of problem it is (algebra, calculus, geometry, trigonometry, etc.)
2. List any relevant formulas or theorems that will be used
3. Show each step clearly with explanations of WHY each step is taken
4. Use proper mathematical notation
5. Provide the final answer clearly marked
6. If applicable, verify the answer or explain how to check it

You MUST respond with ONLY valid JSON (no markdown, no code blocks) in this exact format:
{
    "problem_type": "The category of math problem",
    "concepts": ["List of mathematical concepts used"],
    "steps": [
        {
            "step_number": 1,
            "action": "What is being done in this step",
            "explanation": "Why this step is necessary",
            "result": "The mathematical result of this step"
        }
    ],
    "final_answer": "The final answer to the problem",
    "verification": "How to verify the answer (if applicable)"
}

Always be encouraging and educational. Remember, the goal is to help students LEARN, not just get answers."""

# System prompt for solving problems from images/files
FILE_SOLVER_SYSTEM_PROMPT = """You are a math tutor. Look at this image and solve any math problem you see.

IMPORTANT: Respond with ONLY a JSON object (no markdown, no other text):

{
    "problem_detected": "the math problem from the image",
    "problem_type": "algebra/geometry/calculus/etc",
    "concepts": ["concept1", "concept2"],
    "steps": [
        {"step_number": 1, "action": "What we do", "explanation": "Why", "result": "Result"},
        {"step_number": 2, "action": "Next step", "explanation": "Why", "result": "Result"}
    ],
    "final_answer": "THE FINAL ANSWER",
    "verification": "How to verify"
}

Include at least 2 detailed steps. The final_answer field must contain the actual numerical or algebraic answer."""

# System prompt specifically for PDF/DOCX text-based problems - SIMPLE AND DIRECT
PDF_TEXT_SOLVER_PROMPT = """You are a math tutor. Solve the math problem below step by step.

IMPORTANT: Your response must be ONLY a JSON object with this exact structure (no other text):

{
    "problem_detected": "restate the problem here",
    "problem_type": "algebra",
    "concepts": ["concept1"],
    "steps": [
        {"step_number": 1, "action": "First step", "explanation": "Why we do this", "result": "result"},
        {"step_number": 2, "action": "Second step", "explanation": "Why we do this", "result": "result"}
    ],
    "final_answer": "THE ANSWER",
    "verification": "How to check"
}

Solve this problem:"""

# System prompt for generating quiz questions
QUIZ_SYSTEM_PROMPT = """You are an expert math tutor creating practice problems for students.
Generate quiz questions that test understanding of mathematical concepts.

When creating quiz questions:
1. Create problems that are educational and appropriately challenging
2. Include a mix of difficulty levels when multiple questions are requested
3. Ensure problems are solvable and have clear, unique answers
4. Cover the requested topic area thoroughly

You MUST respond with ONLY valid JSON (no markdown, no code blocks) in this exact format:
{
    "quiz_topic": "The mathematical topic being tested",
    "questions": [
        {
            "question_number": 1,
            "question": "The math problem to solve",
            "difficulty": "easy/medium/hard",
            "hint": "A helpful hint without giving away the answer",
            "correct_answer": "The correct answer",
            "solution_steps": ["Brief steps to solve"]
        }
    ]
}

Create engaging problems that help students build confidence and understanding."""

# System prompt for evaluating quiz answers
EVALUATOR_SYSTEM_PROMPT = """You are an expert math tutor evaluating a student's answer to a math problem.

Compare the student's answer to the correct answer and provide feedback.

You MUST respond with ONLY valid JSON (no markdown, no code blocks) in this exact format:
{
    "is_correct": true or false,
    "feedback": "Encouraging feedback explaining if correct or what went wrong",
    "explanation": "Brief explanation of the correct approach if the answer was wrong"
}

Always be encouraging and constructive, even when the answer is incorrect.
Consider equivalent forms of answers (e.g., 0.5 = 1/2 = 50%)."""

# =============================================================================
# STUDY MODE SYSTEM PROMPTS
# =============================================================================

# System prompt for starting a study session - breaks problem into guided steps
STUDY_START_PROMPT = """You are an expert math tutor helping a student learn by guiding them through a problem step-by-step.
Your role is NOT to solve the problem for them, but to break it down into manageable steps they can work through.

For the given problem:
1. Identify the type of problem and key concepts needed
2. Break down the solution into 3-6 clear steps the student needs to complete
3. For each step, provide a description of what needs to be done (but NOT the answer)
4. Include what mathematical operation or concept is required

You MUST respond with ONLY valid JSON (no markdown, no code blocks) in this exact format:
{
    "problem": "The original problem restated clearly",
    "problem_type": "The category of math problem",
    "concepts_needed": ["List of mathematical concepts the student should know"],
    "total_steps": 4,
    "steps": [
        {
            "step_number": 1,
            "objective": "What the student needs to accomplish in this step",
            "instruction": "Clear instruction on what to do (without giving the answer)",
            "skill_required": "The mathematical skill needed (e.g., 'addition', 'factoring', 'differentiation')",
            "expected_format": "What form the answer should be in (e.g., 'a number', 'an equation', 'simplified expression')"
        }
    ],
    "encouragement": "A brief encouraging message to start the student off"
}

Be encouraging and supportive. Guide them to discover the answer themselves!"""

# System prompt for providing hints during study mode
STUDY_HINT_PROMPT = """You are a supportive math tutor providing a hint to help a student who is stuck.

The student is working on a specific step of a math problem and needs guidance.
Provide a helpful hint that guides them toward the answer WITHOUT giving it away directly.

Rules for good hints:
1. Start with the most gentle hint (a reminder of the concept)
2. If they need more help, provide a more specific hint
3. Never directly state the answer
4. Be encouraging and supportive
5. Relate the hint to concepts they should know

Based on the hint level requested (1 = gentle, 2 = moderate, 3 = strong):
- Level 1: Remind them of the relevant concept or formula
- Level 2: Give a more specific direction or partial setup
- Level 3: Walk them most of the way there, leaving only the final calculation

You MUST respond with ONLY valid JSON (no markdown, no code blocks) in this exact format:
{
    "hint": "The helpful hint text",
    "hint_level": 1,
    "concept_reminder": "A brief reminder of the relevant concept or formula",
    "encouragement": "An encouraging message",
    "next_hint_available": true
}"""

# System prompt for checking a student's step answer in study mode
STUDY_CHECK_PROMPT = """You are a supportive math tutor checking a student's work on a specific step of a problem.

Evaluate if their answer for this step is correct or on the right track.
Be encouraging even if they made a mistake - focus on helping them learn.

Consider:
1. Is the answer mathematically correct for this step?
2. Is it in an acceptable form (equivalent forms should be accepted)?
3. If wrong, what might have caused the error?

You MUST respond with ONLY valid JSON (no markdown, no code blocks) in this exact format:
{
    "is_correct": true or false,
    "feedback": "Specific feedback about their answer",
    "correct_answer": "The correct answer for this step (only if they got it wrong)",
    "error_type": "Type of error if applicable (e.g., 'calculation error', 'wrong operation', 'sign error')",
    "encouragement": "An encouraging message regardless of correctness",
    "tip": "A helpful tip for similar problems in the future"
}

Always be positive and constructive!"""


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def get_api_key_from_request():
    """
    Extract the Gemini API key from the request headers.
    
    The frontend sends the API key in the 'X-API-Key' header.
    This allows each user to use their own API key.
    
    Returns:
        str: The API key from the request header, or None if not provided
    """
    return request.headers.get('X-API-Key')


def clean_json_response(text):
    """
    Clean the response text to extract valid JSON.
    
    Gemini sometimes wraps JSON in markdown code blocks or adds extra text.
    This function extracts just the JSON portion.
    
    Args:
        text: Raw response text from Gemini
        
    Returns:
        Cleaned JSON string
    """
    if not text:
        return "{}"
    
    # Remove markdown code blocks if present
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    text = re.sub(r'```', '', text)
    
    # Remove any leading/trailing whitespace
    text = text.strip()
    
    # Try to find JSON object in the text
    start = text.find('{')
    end = text.rfind('}')
    
    if start != -1 and end != -1:
        text = text[start:end + 1]
    else:
        # If no JSON found, return a default error response
        return '{"error": "No valid JSON found in response"}'
    
    return text.strip()


def allowed_file(filename, file_type='image'):
    """
    Check if a file has an allowed extension.
    
    Args:
        filename: Name of the uploaded file
        file_type: Type of file ('image' or 'document')
        
    Returns:
        Boolean indicating if the file extension is allowed
    """
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    if file_type == 'image':
        return ext in ALLOWED_IMAGE_EXTENSIONS
    elif file_type == 'document':
        return ext in ALLOWED_DOCUMENT_EXTENSIONS
    else:
        return ext in ALLOWED_IMAGE_EXTENSIONS or ext in ALLOWED_DOCUMENT_EXTENSIONS


def get_file_extension(filename):
    """
    Get the file extension from a filename.
    
    Args:
        filename: Name of the file
        
    Returns:
        Lowercase file extension without the dot
    """
    if '.' in filename:
        return filename.rsplit('.', 1)[1].lower()
    return ''


def process_image_file(file_data, filename):
    """
    Process an uploaded image file for Gemini API.
    
    Args:
        file_data: Raw file bytes
        filename: Original filename
        
    Returns:
        PIL Image object ready for Gemini
    """
    image = Image.open(io.BytesIO(file_data))
    
    # Convert to RGB if necessary (handles PNG with transparency, etc.)
    if image.mode in ('RGBA', 'LA', 'P'):
        # Create a white background
        background = Image.new('RGB', image.size, (255, 255, 255))
        if image.mode == 'P':
            image = image.convert('RGBA')
        background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
        image = background
    elif image.mode != 'RGB':
        image = image.convert('RGB')
    
    return image


def extract_text_from_pdf(file_data):
    """
    Extract text content from a PDF file.
    
    Args:
        file_data: Raw PDF file bytes
        
    Returns:
        Tuple of (extracted_text, list_of_page_images)
    """
    text_content = ""
    page_images = []
    
    if PYMUPDF_AVAILABLE:
        # Use PyMuPDF for text extraction and image conversion
        pdf_document = fitz.open(stream=file_data, filetype="pdf")
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Extract text
            text_content += f"\n--- Page {page_num + 1} ---\n"
            text_content += page.get_text()
            
            # Convert page to image (for visual math problems)
            mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            page_images.append(img)
        
        pdf_document.close()
    
    elif PDF2IMAGE_AVAILABLE:
        # Fallback to pdf2image
        try:
            images = convert_from_bytes(file_data, dpi=150)
            page_images = images
            text_content = "PDF converted to images for visual analysis."
        except Exception as e:
            text_content = f"Error converting PDF: {str(e)}"
    
    else:
        text_content = "PDF processing libraries not available. Please install PyMuPDF (fitz) or pdf2image."
    
    return text_content, page_images


def extract_text_from_docx(file_data):
    """
    Extract text content from a DOCX file.
    
    Args:
        file_data: Raw DOCX file bytes
        
    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        return "DOCX processing library not available. Please install python-docx."
    
    try:
        doc = DocxDocument(io.BytesIO(file_data))
        
        text_content = ""
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text_content += " | ".join(row_text) + "\n"
        
        return text_content.strip()
    
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"


def call_gemini(prompt, system_prompt, api_key):
    """
    Make a request to the Gemini API using the user's API key.
    
    Args:
        prompt: The user's prompt/question
        system_prompt: Instructions for how Gemini should respond
        api_key: The user's Gemini API key
        
    Returns:
        Parsed JSON response from Gemini
        
    Raises:
        ValueError: If API key is not provided
        Exception: If API call fails
    """
    if not api_key:
        raise ValueError("API key is required. Please provide your Gemini API key.")
    
    # Configure the Gemini API with the user's key
    genai.configure(api_key=api_key)
    
    # Initialize the Gemini model
    model = genai.GenerativeModel('gemini-2.0-flash')
    
    # Combine system prompt with user prompt
    full_prompt = f"{system_prompt}\n\n{prompt}"
    
    # Generate response from Gemini
    response = model.generate_content(full_prompt)
    
    # Extract text from response
    response_text = response.text
    
    # Clean and parse JSON
    cleaned_text = clean_json_response(response_text)
    
    try:
        return json.loads(cleaned_text)
    except json.JSONDecodeError as e:
        # If JSON parsing fails, try to construct a response from the text
        print(f"JSON parse error in call_gemini: {e}")
        print(f"Raw response (first 500 chars): {response_text[:500]}")
        
        # Try to extract useful information from the response
        final_answer = ""
        if "answer" in response_text.lower():
            lines = response_text.split('\n')
            for line in lines:
                if "answer" in line.lower() and ('=' in line or ':' in line):
                    final_answer = line.split('=')[-1].split(':')[-1].strip()
                    break
        
        return {
            "problem_detected": "Problem analyzed",
            "problem_type": "Mathematical Problem",
            "concepts": [],
            "steps": [
                {
                    "step_number": 1,
                    "action": "Solution",
                    "explanation": response_text[:1000] if len(response_text) > 1000 else response_text,
                    "result": final_answer or "See explanation"
                }
            ],
            "final_answer": final_answer or "See the explanation above for the complete solution",
            "verification": "Verify by checking the steps above"
        }


def call_gemini_with_image(images, prompt, system_prompt, api_key):
    """
    Make a request to the Gemini API with image(s) using the user's API key.
    
    Args:
        images: List of PIL Image objects or single PIL Image
        prompt: Additional text prompt/question
        system_prompt: Instructions for how Gemini should respond
        api_key: The user's Gemini API key
        
    Returns:
        Parsed JSON response from Gemini
        
    Raises:
        ValueError: If API key is not provided
        Exception: If API call fails
    """
    if not api_key:
        raise ValueError("API key is required. Please provide your Gemini API key.")
    
    # Configure the Gemini API with the user's key
    genai.configure(api_key=api_key)
    
    # Initialize the Gemini model with vision capabilities
    model = genai.GenerativeModel('gemini-2.0-flash')
    
    # Prepare content parts - images first, then text
    content_parts = []
    
    # Add images first (Gemini works better with images before text)
    if isinstance(images, list):
        for img in images[:3]:  # Limit to first 3 images for reliability
            content_parts.append(img)
    else:
        content_parts.append(images)
    
    # Add system prompt and user prompt as text
    full_prompt = system_prompt
    if prompt:
        full_prompt += f"\n\nAdditional context from user: {prompt}"
    
    content_parts.append(full_prompt)
    
    # Generate response from Gemini
    try:
        response = model.generate_content(content_parts)
        
        # Extract text from response
        response_text = response.text
        
        # Clean and parse JSON
        cleaned_text = clean_json_response(response_text)
        
        try:
            return json.loads(cleaned_text)
        except json.JSONDecodeError as e:
            print(f"JSON parse error: {e}")
            print(f"Raw response: {response_text[:500]}...")
            
            return {
                "problem_detected": "Problem analyzed from uploaded file",
                "problem_type": "Mathematical Problem",
                "concepts": [],
                "steps": [
                    {
                        "step_number": 1,
                        "action": "Analysis",
                        "explanation": response_text,
                        "result": "See explanation above"
                    }
                ],
                "final_answer": "Please see the step-by-step explanation above",
                "verification": "N/A"
            }
    except Exception as e:
        print(f"Gemini API error: {e}")
        raise


def validate_solution_response(solution, source_description="uploaded file"):
    """
    Validate and fill in missing fields in a solution response.
    
    This ensures the frontend always receives a complete response structure
    even if the AI didn't return all expected fields.
    
    Args:
        solution: The solution dictionary from Gemini
        source_description: Description of where the problem came from
        
    Returns:
        Validated solution dictionary with all required fields
    """
    if not solution:
        solution = {}
    
    # List of invalid/placeholder answers to detect
    invalid_answers = ['N/A', 'n/a', '', None, 'Unable to determine', 
                       'See steps above', 'No math problem found', 
                       'Could not identify', 'Unable to solve']
    
    def is_valid_answer(ans):
        if not ans:
            return False
        ans_str = str(ans).strip().lower()
        for invalid in invalid_answers:
            if invalid and invalid.lower() in ans_str:
                return False
        return len(ans_str) > 0
    
    # Ensure problem_detected field exists
    if not solution.get('problem_detected'):
        solution['problem_detected'] = f"Problem from {source_description}"
    
    # Ensure problem_type field exists
    if not solution.get('problem_type') or solution.get('problem_type') == 'N/A':
        solution['problem_type'] = "Mathematical Problem"
    
    # Ensure concepts is a list
    if not solution.get('concepts') or not isinstance(solution.get('concepts'), list):
        solution['concepts'] = []
    
    # Check if we have valid steps with real results
    has_valid_steps = False
    if solution.get('steps') and isinstance(solution.get('steps'), list) and len(solution['steps']) > 0:
        for step in solution['steps']:
            if isinstance(step, dict) and step.get('result') and is_valid_answer(step.get('result')):
                has_valid_steps = True
                break
    
    # Ensure steps is a non-empty list with proper structure
    if not has_valid_steps:
        if solution.get('final_answer') and is_valid_answer(solution.get('final_answer')):
            solution['steps'] = [
                {
                    "step_number": 1,
                    "action": "Solution",
                    "explanation": "The problem was analyzed and solved.",
                    "result": str(solution.get('final_answer', ''))
                }
            ]
        else:
            solution['steps'] = [
                {
                    "step_number": 1,
                    "action": "Processing Issue",
                    "explanation": f"The AI could not generate a proper solution. Extracted content: {source_description[:200] if len(source_description) > 50 else source_description}",
                    "result": "Please try typing the problem manually"
                }
            ]
    else:
        # Validate each step has required fields
        for i, step in enumerate(solution['steps']):
            if not isinstance(step, dict):
                solution['steps'][i] = {
                    "step_number": i + 1,
                    "action": "Step",
                    "explanation": str(step),
                    "result": ""
                }
            else:
                step['step_number'] = step.get('step_number', i + 1)
                step['action'] = step.get('action', 'Step')
                step['explanation'] = step.get('explanation', '')
                result = step.get('result', '')
                step['result'] = result if is_valid_answer(result) else ''
    
    # Ensure final_answer exists and is valid
    if not is_valid_answer(solution.get('final_answer')):
        if solution['steps'] and is_valid_answer(solution['steps'][-1].get('result')):
            solution['final_answer'] = solution['steps'][-1]['result']
        else:
            solution['final_answer'] = "Unable to determine - please try entering the problem as text"
    
    # Ensure verification exists
    if not solution.get('verification') or solution.get('verification') == 'N/A':
        solution['verification'] = "Verify by substituting the answer back into the original problem"
    
    return solution


# =============================================================================
# API ROUTES
# =============================================================================

@app.route('/')
def serve_frontend():
    """
    Serve the main frontend HTML file.
    
    This allows the entire application to be accessed from a single URL:
    http://localhost:5000
    
    Returns:
        The index.html file containing the React frontend
    """
    return send_from_directory('.', 'index.html')


@app.route('/api/health', methods=['GET'])
def health_check():
    """
    Health check endpoint to verify the server is running.
    
    Returns:
        JSON response with status "healthy" and HTTP 200
    """
    return jsonify({
        "status": "healthy", 
        "message": "AI Math Tutor server is running",
        "model": "Google Gemini 2.0 Flash",
        "auth": "Google Sign-In with user-provided API keys",
        "features": {
            "file_upload": True,
            "study_mode": True,
            "supported_formats": list(ALLOWED_IMAGE_EXTENSIONS) + list(ALLOWED_DOCUMENT_EXTENSIONS),
            "pymupdf_available": PYMUPDF_AVAILABLE,
            "pdf2image_available": PDF2IMAGE_AVAILABLE,
            "docx_available": DOCX_AVAILABLE
        }
    })


@app.route('/api/config', methods=['GET'])
def get_config():
    """
    Get frontend configuration including Google Client ID and supported file types.
    
    Returns:
        JSON response with configuration values
    """
    return jsonify({
        "google_client_id": os.getenv('GOOGLE_CLIENT_ID', ''),
        "use_google_auth": bool(os.getenv('GOOGLE_CLIENT_ID')),
        "supported_file_types": {
            "images": list(ALLOWED_IMAGE_EXTENSIONS),
            "documents": list(ALLOWED_DOCUMENT_EXTENSIONS)
        },
        "max_file_size_mb": 16
    })


@app.route('/api/verify-key', methods=['POST'])
def verify_api_key():
    """
    Verify that a user's Gemini API key is valid.
    
    Request Headers:
        X-API-Key: The Gemini API key to verify
    
    Returns:
        JSON response indicating if the key is valid
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "valid": False,
                "error": "API key is required"
            }), 400
        
        # Configure and test the API key with a simple request
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Make a simple test request
        response = model.generate_content("Reply with just the word 'OK'")
        
        return jsonify({
            "valid": True,
            "message": "API key is valid!"
        })
        
    except Exception as e:
        error_message = str(e)
        
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower() or "401" in error_message:
            return jsonify({
                "valid": False,
                "error": "Invalid API key. Please check your key and try again."
            }), 401
        
        return jsonify({
            "valid": False,
            "error": f"Error verifying API key: {error_message}"
        }), 500


@app.route('/api/solve', methods=['POST'])
def solve_problem():
    """
    Solve a math problem with step-by-step explanations.
    
    Request Headers:
        X-API-Key: The user's Gemini API key
    
    Request Body (JSON):
        {
            "problem": "The math problem to solve (string)"
        }
    
    Returns:
        JSON response containing step-by-step solution
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "error": "API key is required. Please sign in and provide your Gemini API key.",
                "code": "NO_API_KEY"
            }), 401
        
        data = request.get_json()
        
        if not data or 'problem' not in data:
            return jsonify({
                "error": "Missing 'problem' in request body",
                "example": {"problem": "Solve for x: 2x + 5 = 13"}
            }), 400
        
        problem = data['problem']
        
        if not problem.strip():
            return jsonify({"error": "Problem cannot be empty"}), 400
        
        solution = call_gemini(
            f"Please solve this math problem step-by-step:\n\n{problem}",
            SOLVER_SYSTEM_PROMPT,
            api_key
        )
        
        return jsonify(solution)
        
    except json.JSONDecodeError as je:
        return jsonify({
            "error": "Failed to parse AI response. Please try again.",
            "details": str(je)
        }), 500
        
    except ValueError as ve:
        return jsonify({
            "error": str(ve),
            "code": "API_KEY_ERROR"
        }), 401
        
    except Exception as e:
        error_message = str(e)
        
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower() or "401" in error_message:
            return jsonify({
                "error": "Invalid API key. Please check your Gemini API key.",
                "code": "INVALID_API_KEY",
                "help": "Get a free key at: https://aistudio.google.com/apikey"
            }), 401
            
        return jsonify({"error": f"Server Error: {error_message}"}), 500


@app.route('/api/solve/file', methods=['POST'])
def solve_from_file():
    """
    Solve a math problem from an uploaded file (image, PDF, or DOCX).
    
    Request Headers:
        X-API-Key: The user's Gemini API key
    
    Request Form Data:
        file: The uploaded file (image, PDF, or DOCX)
        additional_context: Optional additional context or question about the problem
    
    Returns:
        JSON response containing step-by-step solution
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "error": "API key is required. Please sign in and provide your Gemini API key.",
                "code": "NO_API_KEY"
            }), 401
        
        if 'file' not in request.files:
            return jsonify({
                "error": "No file uploaded.",
                "supported_types": list(ALLOWED_IMAGE_EXTENSIONS) + list(ALLOWED_DOCUMENT_EXTENSIONS)
            }), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        additional_context = request.form.get('additional_context', '')
        file_ext = get_file_extension(file.filename)
        
        if not allowed_file(file.filename, 'all'):
            return jsonify({
                "error": f"Unsupported file type: .{file_ext}",
                "supported_types": list(ALLOWED_IMAGE_EXTENSIONS) + list(ALLOWED_DOCUMENT_EXTENSIONS)
            }), 415
        
        file_data = file.read()
        
        # Process based on file type
        if file_ext in ALLOWED_IMAGE_EXTENSIONS:
            try:
                image = process_image_file(file_data, file.filename)
                solution = call_gemini_with_image(
                    image,
                    additional_context or "Please solve the math problem shown in this image.",
                    FILE_SOLVER_SYSTEM_PROMPT,
                    api_key
                )
                solution = validate_solution_response(solution, f"image: {file.filename}")
            except Exception as img_error:
                return jsonify({"error": f"Failed to process image: {str(img_error)}"}), 400
        
        elif file_ext == 'pdf':
            text_content, page_images = extract_text_from_pdf(file_data)
            
            if text_content:
                text_content = re.sub(r'---\s*Page\s*\d+\s*---', '\n', text_content)
                text_content = re.sub(r'\n{3,}', '\n\n', text_content)
                text_content = text_content.strip()
            
            solution = None
            
            if text_content and len(text_content.strip()) > 10:
                try:
                    pdf_prompt = f"{text_content}\n\n{additional_context}" if additional_context else text_content
                    solution = call_gemini(pdf_prompt, PDF_TEXT_SOLVER_PROMPT, api_key)
                    
                    if solution and solution.get('final_answer') and solution['final_answer'] not in ['N/A', '', 'See steps above for the solution']:
                        solution = validate_solution_response(solution, "PDF text analysis")
                    else:
                        solution = None
                except Exception:
                    solution = None
            
            if not solution and page_images:
                try:
                    img_prompt = "Solve the math problem shown in this image step by step."
                    if text_content:
                        img_prompt += f" The text reads: {text_content[:300]}"
                    if additional_context:
                        img_prompt += f" {additional_context}"
                    
                    solution = call_gemini_with_image(page_images[0], img_prompt, FILE_SOLVER_SYSTEM_PROMPT, api_key)
                    solution = validate_solution_response(solution, "PDF image analysis")
                except Exception:
                    pass
            
            if not solution:
                return jsonify({
                    "error": "Could not solve the problem from this PDF.",
                    "hint": "Try typing the problem manually or uploading a clearer image."
                }), 400
        
        elif file_ext in ['docx', 'doc']:
            text_content = extract_text_from_docx(file_data)
            
            if text_content and not text_content.startswith("Error"):
                docx_prompt = f"Document Content:\n{text_content}\n\n{f'Additional context: {additional_context}' if additional_context else ''}"
                solution = call_gemini(docx_prompt, PDF_TEXT_SOLVER_PROMPT, api_key)
                solution = validate_solution_response(solution, f"Word document: {file.filename}")
            else:
                return jsonify({
                    "error": text_content if text_content else "Could not extract content from DOCX file."
                }), 400
        else:
            return jsonify({"error": "Unsupported file type"}), 415
        
        solution['source_file'] = {
            'filename': file.filename,
            'type': file_ext,
            'size_bytes': len(file_data)
        }
        
        return jsonify(solution)
        
    except ValueError as ve:
        return jsonify({"error": str(ve), "code": "API_KEY_ERROR"}), 401
        
    except Exception as e:
        error_message = str(e)
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower():
            return jsonify({
                "error": "Invalid API key. Please check your Gemini API key.",
                "code": "INVALID_API_KEY"
            }), 401
        return jsonify({"error": f"Server Error: {error_message}"}), 500


# =============================================================================
# STUDY MODE API ROUTES
# =============================================================================

@app.route('/api/study/start', methods=['POST'])
def start_study_session():
    """
    Start a new study session for a math problem.
    
    This endpoint analyzes a problem and breaks it down into guided steps
    that the student will work through interactively.
    
    Request Headers:
        X-API-Key: The user's Gemini API key
    
    Request Body (JSON):
        {
            "problem": "The math problem to study"
        }
    
    Returns:
        JSON response containing:
        - problem: The original problem
        - problem_type: Category of the math problem
        - concepts_needed: List of concepts the student should know
        - total_steps: Number of steps to complete
        - steps: Array of step objects with objectives and instructions
        - encouragement: Encouraging message to start
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "error": "API key is required. Please sign in and provide your Gemini API key.",
                "code": "NO_API_KEY"
            }), 401
        
        data = request.get_json()
        
        if not data or 'problem' not in data:
            return jsonify({
                "error": "Missing 'problem' in request body",
                "example": {"problem": "Solve for x: 2x + 5 = 13"}
            }), 400
        
        problem = data['problem']
        
        if not problem.strip():
            return jsonify({"error": "Problem cannot be empty"}), 400
        
        # Call Gemini to break down the problem into study steps
        study_plan = call_gemini(
            f"Please analyze this math problem and create a guided study plan:\n\n{problem}",
            STUDY_START_PROMPT,
            api_key
        )
        
        # Validate the response has required fields
        if not study_plan.get('steps'):
            study_plan['steps'] = [{
                "step_number": 1,
                "objective": "Solve the problem",
                "instruction": "Work through the problem step by step",
                "skill_required": "Mathematical reasoning",
                "expected_format": "The final answer"
            }]
        
        if not study_plan.get('total_steps'):
            study_plan['total_steps'] = len(study_plan['steps'])
        
        if not study_plan.get('problem'):
            study_plan['problem'] = problem
        
        if not study_plan.get('encouragement'):
            study_plan['encouragement'] = "Let's work through this problem together! Take your time with each step."
        
        return jsonify(study_plan)
        
    except json.JSONDecodeError as je:
        return jsonify({
            "error": "Failed to parse AI response. Please try again.",
            "details": str(je)
        }), 500
        
    except ValueError as ve:
        return jsonify({
            "error": str(ve),
            "code": "API_KEY_ERROR"
        }), 401
        
    except Exception as e:
        error_message = str(e)
        
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower() or "401" in error_message:
            return jsonify({
                "error": "Invalid API key. Please check your Gemini API key.",
                "code": "INVALID_API_KEY",
                "help": "Get a free key at: https://aistudio.google.com/apikey"
            }), 401
            
        return jsonify({"error": f"Server Error: {error_message}"}), 500


@app.route('/api/study/hint', methods=['POST'])
def get_study_hint():
    """
    Get a hint for a specific step in study mode.
    
    Request Headers:
        X-API-Key: The user's Gemini API key
    
    Request Body (JSON):
        {
            "problem": "The original math problem",
            "step_number": 1,
            "step_objective": "What the student needs to do",
            "hint_level": 1 (1=gentle, 2=moderate, 3=strong),
            "student_attempt": "What the student has tried so far (optional)"
        }
    
    Returns:
        JSON response containing:
        - hint: The hint text
        - hint_level: The level of hint provided
        - concept_reminder: A reminder of the relevant concept
        - encouragement: An encouraging message
        - next_hint_available: Whether a stronger hint is available
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "error": "API key is required.",
                "code": "NO_API_KEY"
            }), 401
        
        data = request.get_json()
        
        required_fields = ['problem', 'step_number', 'step_objective']
        for field in required_fields:
            if field not in data:
                return jsonify({
                    "error": f"Missing '{field}' in request body",
                    "required_fields": required_fields
                }), 400
        
        problem = data['problem']
        step_number = data['step_number']
        step_objective = data['step_objective']
        hint_level = data.get('hint_level', 1)
        student_attempt = data.get('student_attempt', '')
        
        # Build the hint request prompt
        hint_prompt = f"""Problem: {problem}

Current Step: Step {step_number}
Step Objective: {step_objective}

Hint Level Requested: {hint_level} (1=gentle reminder, 2=more specific guidance, 3=strong hint)

{"Student's attempt so far: " + student_attempt if student_attempt else "Student hasn't attempted yet."}

Please provide an appropriate hint for this level."""
        
        hint_response = call_gemini(hint_prompt, STUDY_HINT_PROMPT, api_key)
        
        # Validate response
        if not hint_response.get('hint'):
            hint_response['hint'] = "Think about what mathematical operation would help you here."
        
        hint_response['hint_level'] = hint_level
        hint_response['next_hint_available'] = hint_level < 3
        
        if not hint_response.get('encouragement'):
            hint_response['encouragement'] = "You're doing great! Keep thinking through it."
        
        return jsonify(hint_response)
        
    except Exception as e:
        error_message = str(e)
        
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower():
            return jsonify({
                "error": "Invalid API key.",
                "code": "INVALID_API_KEY"
            }), 401
            
        return jsonify({"error": f"Server Error: {error_message}"}), 500


@app.route('/api/study/check', methods=['POST'])
def check_study_step():
    """
    Check a student's answer for a specific step in study mode.
    
    Request Headers:
        X-API-Key: The user's Gemini API key
    
    Request Body (JSON):
        {
            "problem": "The original math problem",
            "step_number": 1,
            "step_objective": "What the student needed to do",
            "student_answer": "The student's answer for this step",
            "expected_format": "What form the answer should be in"
        }
    
    Returns:
        JSON response containing:
        - is_correct: Boolean indicating if the answer is correct
        - feedback: Specific feedback about the answer
        - correct_answer: The correct answer (only if wrong)
        - error_type: Type of error if applicable
        - encouragement: An encouraging message
        - tip: A helpful tip for the future
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "error": "API key is required.",
                "code": "NO_API_KEY"
            }), 401
        
        data = request.get_json()
        
        required_fields = ['problem', 'step_number', 'step_objective', 'student_answer']
        for field in required_fields:
            if field not in data:
                return jsonify({
                    "error": f"Missing '{field}' in request body",
                    "required_fields": required_fields
                }), 400
        
        problem = data['problem']
        step_number = data['step_number']
        step_objective = data['step_objective']
        student_answer = data['student_answer']
        expected_format = data.get('expected_format', 'an answer')
        
        # Build the check prompt
        check_prompt = f"""Original Problem: {problem}

Step Being Checked: Step {step_number}
Step Objective: {step_objective}
Expected Answer Format: {expected_format}

Student's Answer: {student_answer}

Please evaluate if the student's answer is correct for this step. Consider equivalent forms."""
        
        check_response = call_gemini(check_prompt, STUDY_CHECK_PROMPT, api_key)
        
        # Validate response
        if 'is_correct' not in check_response:
            check_response['is_correct'] = False
        
        if not check_response.get('feedback'):
            check_response['feedback'] = "Let me check your work..."
        
        if not check_response.get('encouragement'):
            if check_response['is_correct']:
                check_response['encouragement'] = "Excellent work! You got it right!"
            else:
                check_response['encouragement'] = "Good effort! Let's look at what happened."
        
        return jsonify(check_response)
        
    except Exception as e:
        error_message = str(e)
        
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower():
            return jsonify({
                "error": "Invalid API key.",
                "code": "INVALID_API_KEY"
            }), 401
            
        return jsonify({"error": f"Server Error: {error_message}"}), 500


@app.route('/api/study/solution', methods=['POST'])
def get_step_solution():
    """
    Get the full solution for a specific step (when student gives up or wants to see answer).
    
    Request Headers:
        X-API-Key: The user's Gemini API key
    
    Request Body (JSON):
        {
            "problem": "The original math problem",
            "step_number": 1,
            "step_objective": "What the student needed to do"
        }
    
    Returns:
        JSON response containing the solution for that step
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "error": "API key is required.",
                "code": "NO_API_KEY"
            }), 401
        
        data = request.get_json()
        
        required_fields = ['problem', 'step_number', 'step_objective']
        for field in required_fields:
            if field not in data:
                return jsonify({
                    "error": f"Missing '{field}' in request body"
                }), 400
        
        problem = data['problem']
        step_number = data['step_number']
        step_objective = data['step_objective']
        
        # Build the solution prompt
        solution_prompt = f"""Problem: {problem}

The student needs help with Step {step_number}: {step_objective}

Please provide the solution for ONLY this specific step with a clear explanation.

Respond with ONLY valid JSON in this format:
{{
    "step_solution": "The answer for this step",
    "explanation": "Clear explanation of how to get this answer",
    "key_concept": "The main mathematical concept used",
    "tip": "A tip for solving similar steps in the future"
}}"""
        
        system_prompt = "You are a helpful math tutor. Provide the solution for the requested step. Respond with ONLY valid JSON."
        
        solution = call_gemini(solution_prompt, system_prompt, api_key)
        
        if not solution.get('step_solution'):
            solution['step_solution'] = "Please see the explanation."
        
        if not solution.get('explanation'):
            solution['explanation'] = "Work through the problem using the objective given."
        
        return jsonify(solution)
        
    except Exception as e:
        error_message = str(e)
        
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower():
            return jsonify({
                "error": "Invalid API key.",
                "code": "INVALID_API_KEY"
            }), 401
            
        return jsonify({"error": f"Server Error: {error_message}"}), 500


# =============================================================================
# QUIZ API ROUTES
# =============================================================================

@app.route('/api/quiz/generate', methods=['POST'])
def generate_quiz():
    """
    Generate quiz questions for a specified math topic.
    
    Request Headers:
        X-API-Key: The user's Gemini API key
    
    Request Body (JSON):
        {
            "topic": "Math topic (e.g., 'algebra', 'calculus', 'trigonometry')",
            "num_questions": Number of questions to generate (default: 3),
            "difficulty": "easy", "medium", "hard", or "mixed" (default: "mixed")
        }
    
    Returns:
        JSON response containing quiz questions
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "error": "API key is required. Please sign in and provide your Gemini API key.",
                "code": "NO_API_KEY"
            }), 401
        
        data = request.get_json()
        
        if not data or 'topic' not in data:
            return jsonify({
                "error": "Missing 'topic' in request body",
                "example": {"topic": "algebra", "num_questions": 3, "difficulty": "medium"}
            }), 400
        
        topic = data['topic']
        num_questions = data.get('num_questions', 3)
        difficulty = data.get('difficulty', 'mixed')
        
        num_questions = min(max(1, num_questions), 10)
        
        quiz = call_gemini(
            f"Generate {num_questions} {difficulty} difficulty quiz questions about {topic}.",
            QUIZ_SYSTEM_PROMPT,
            api_key
        )
        
        return jsonify(quiz)
        
    except json.JSONDecodeError:
        return jsonify({"error": "Failed to generate quiz. Please try again."}), 500
        
    except ValueError as ve:
        return jsonify({"error": str(ve), "code": "API_KEY_ERROR"}), 401
        
    except Exception as e:
        error_message = str(e)
        
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower() or "401" in error_message:
            return jsonify({
                "error": "Invalid API key. Please check your Gemini API key.",
                "code": "INVALID_API_KEY"
            }), 401
            
        return jsonify({"error": f"Server Error: {error_message}"}), 500


@app.route('/api/quiz/evaluate', methods=['POST'])
def evaluate_answer():
    """
    Evaluate a student's answer to a quiz question.
    
    Request Headers:
        X-API-Key: The user's Gemini API key
    
    Request Body (JSON):
        {
            "question": "The quiz question that was asked",
            "correct_answer": "The correct answer",
            "student_answer": "The student's submitted answer"
        }
    
    Returns:
        JSON response containing evaluation feedback
    """
    try:
        api_key = get_api_key_from_request()
        
        if not api_key:
            return jsonify({
                "error": "API key is required. Please sign in and provide your Gemini API key.",
                "code": "NO_API_KEY"
            }), 401
        
        data = request.get_json()
        
        required_fields = ['question', 'correct_answer', 'student_answer']
        for field in required_fields:
            if field not in data:
                return jsonify({
                    "error": f"Missing '{field}' in request body",
                    "required_fields": required_fields
                }), 400
        
        question = data['question']
        correct_answer = data['correct_answer']
        student_answer = data['student_answer']
        
        evaluation = call_gemini(
            f"""Evaluate this student's answer:
            
Question: {question}
Correct Answer: {correct_answer}
Student's Answer: {student_answer}

Please determine if the student's answer is correct (considering equivalent forms) and provide feedback.""",
            EVALUATOR_SYSTEM_PROMPT,
            api_key
        )
        
        return jsonify(evaluation)
        
    except json.JSONDecodeError:
        is_correct = str(student_answer).strip().lower() == str(correct_answer).strip().lower()
        return jsonify({
            "is_correct": is_correct,
            "feedback": "Correct! Great job!" if is_correct else "Not quite right. Keep practicing!",
            "explanation": "" if is_correct else f"The correct answer was: {correct_answer}"
        })
        
    except ValueError as ve:
        return jsonify({"error": str(ve), "code": "API_KEY_ERROR"}), 401
        
    except Exception as e:
        error_message = str(e)
        
        if "API_KEY" in error_message.upper() or "invalid" in error_message.lower() or "401" in error_message:
            return jsonify({
                "error": "Invalid API key. Please check your Gemini API key.",
                "code": "INVALID_API_KEY"
            }), 401
            
        return jsonify({"error": f"Server Error: {error_message}"}), 500


# =============================================================================
# MAIN ENTRY POINT
# =============================================================================

if __name__ == '__main__':
    """
    Main entry point for the Flask application.
    """
    print("=" * 60)
    print("AI Math Tutor - Backend Server")
    print("=" * 60)
    print()
    print("  Powered by Google Gemini (Users provide their own API key)")
    print()
    print("  Open your browser and go to:")
    print()
    print("     http://localhost:5000")
    print()
    print("  Features:")
    print("     ✓  Google Sign-In authentication")
    print("     ✓  User-provided Gemini API keys")
    print("     ✓  Step-by-step math solutions")
    print("     ✓  Interactive practice quizzes")
    print("     ✓  STUDY MODE - Interactive guided learning")
    print("     ✓  FILE UPLOAD SUPPORT:")
    print(f"        - Images: {', '.join(ALLOWED_IMAGE_EXTENSIONS)}")
    print(f"        - Documents: {', '.join(ALLOWED_DOCUMENT_EXTENSIONS)}")
    print()
    print("  Library Status:")
    print(f"     {'✓' if PYMUPDF_AVAILABLE else '✗'}  PyMuPDF (PDF processing)")
    print(f"     {'✓' if PDF2IMAGE_AVAILABLE else '✗'}  pdf2image (PDF to image)")
    print(f"     {'✓' if DOCX_AVAILABLE else '✗'}  python-docx (Word documents)")
    print()
    print("=" * 60)
    
    app.run(debug=True, host='0.0.0.0', port=5000)